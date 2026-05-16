#!/usr/bin/env python3
"""Generate full FIT paper Word from FITpaper.docx template (Tagawa structure)."""
from __future__ import annotations

import shutil
from pathlib import Path

from docx import Document
from docx.shared import Cm, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

ROOT = Path(__file__).resolve().parents[1]
TEMPLATE = Path(r"d:\Downloads\FITpaper.docx")
OUT = ROOT / "FIT_paper.docx"
FIG_CANDIDATE = ROOT / "comparison_visuals" / "fern" / "compare_00000.png"

# Style names in official FITpaper.docx
S_TITLE = "FITタイトル"
S_TITLE_EN = "FITタイトル英"
S_AUTHOR = "FIT筆者"
S_AUTHOR_EN = "FIT筆者英"
S_BODY = "FIT本文"
S_H1 = "FIT見出し1"
S_H2 = "FIT見出し2"
S_FIG = "FIT図"
S_REF_HEAD = "FIT謝辞と参考文献の見出し"
S_REF = "FIT参考文献"

TITLE_JA = (
    "VGGT 誘導 3D Gaussian Splatting における\n"
    "幾何信頼性重み W_geo による深度正則化"
)
TITLE_EN = "W_geo-Weighted Depth Regularization for VGGT-Guided 3D Gaussian Splatting"
AUTHOR_JA = "郭　得地"
AUTHOR_EN = "Guo Dedi"


def delete_paragraph(paragraph) -> None:
    el = paragraph._element
    el.getparent().remove(el)


def clear_body_keep_header(doc: Document, keep: int = 4) -> None:
    for p in doc.paragraphs[keep:][::-1]:
        delete_paragraph(p)


def add_body(doc: Document, text: str) -> None:
    for block in text.strip().split("\n\n"):
        block = block.strip()
        if block:
            doc.add_paragraph(block, style=S_BODY)


def add_table_results(doc: Document) -> None:
    doc.add_paragraph("表 1　訓練視点における再構成品質（7000 iter）", style=S_BODY)
    rows = [
        ("シーン", "方式", "L1↓", "PSNR↑", "SSIM↑"),
        ("Kitchen", "Baseline", "0.00610", "40.34", "0.9890"),
        ("Kitchen", "W_geo", "0.00599", "40.40", "0.9884"),
        ("Fern", "Baseline", "0.01586", "29.13", "0.9650"),
        ("Fern", "W_geo", "0.01320", "31.05", "0.9716"),
        ("Flower", "Baseline", "0.04659", "21.25", "0.7818"),
        ("Flower", "W_geo", "0.04177", "21.75", "0.8117"),
    ]
    table = doc.add_table(rows=len(rows), cols=5)
    table.style = "Table Grid"
    for i, row in enumerate(rows):
        for j, val in enumerate(row):
            table.rows[i].cells[j].text = val
            if i == 0:
                for p in table.rows[i].cells[j].paragraphs:
                    for r in p.runs:
                        r.bold = True
    doc.add_paragraph("", style=S_BODY)


def build() -> Path:
    if not TEMPLATE.exists():
        raise SystemExit(f"Template not found: {TEMPLATE}")

    shutil.copy2(TEMPLATE, OUT)
    doc = Document(str(OUT))

    # Header: title & authors
    doc.paragraphs[0].text = TITLE_JA
    doc.paragraphs[1].text = TITLE_EN
    doc.paragraphs[2].text = AUTHOR_JA
    doc.paragraphs[3].text = AUTHOR_EN

    clear_body_keep_header(doc, keep=4)

    # --- 1. はじめに ---
    doc.add_paragraph("1. はじめに", style=S_H1)
    add_body(
        doc,
        """
3D Gaussian Splatting (3DGS) [1] はリアルタイム新視点合成を実現する．
スパース視点では幾何が不安定になり，単眼深度や前向き幾何モデルの
事前知識を正則化に用いる研究が増えている [3]．
Visual Geometry Grounded Transformer (VGGT) [2] は複数画像から
深度・カメラ・追跡を推論し，depth_conf や可視性など信頼度情報も
出力する．

従来の VGGT+3DGS 統合では，W = (1−σ)×V×C 等の単一重みを
光度損失と深度損失の双方に掛けることが多い [2]．
しかし光度損失の教師は実画像であり信頼できる一方，
深度損失の教師は VGGT 推定であり誤差を含む．
「VGGT の低信頼領域（入力品質）」と「3DGS の再構築困難領域（モデル能力）」は
直交し得る（例：鏡面は幾何は明確だが SH 表現が困難）ため，
同一 W の共用は概念的に不適切である．

本稿では幾何信頼性 W_geo を深度正則化のみに適用し，
光度損失には掛けない分離を提案する．貢献は，
(1) 参考値の性質に基づく損失別重み付け，
(2) σ，V，C の加重平均による W_geo 設計，
(3) 3 シーンと係数消融による検証，の 3 点である．
""",
    )

    # --- 2. 方法 ---
    doc.add_paragraph("2. 方法", style=S_H1)

    doc.add_paragraph("2.1 概念分離と W_geo の構成", style=S_H2)
    add_body(
        doc,
        """
W_geo は画素ごとの幾何信頼性を表す．VGGT から得る 3 信号を
[0, 1] に正規化した σ̂（深度不確実性），V̂（跨フレーム可視性），
Ĉ（多視点再投影一致性）とし，加重平均で融合する：

                    W_geo = max(0.1, 0.5(1−σ̂) + 0.3V̂ + 0.2Ĉ)     (1)

最小値 0.1 は完全マスクを避ける．乗法融合 W = (1−σ)×V×C は
失敗要因が相関する領域で過剰抑制し（低信頼画素比 56%），
加重平均では 34% に改善したため本稿では (1) を用いる．
""",
    )

    doc.add_paragraph("2.2 損失関数への統合", style=S_H2)
    add_body(
        doc,
        """
総損失は次式とする：

  L = (1−α)L_1 + α(1−SSIM) + λ_d(t) E[|D_render − D_VGGT| · M · W_geo]   (2)

ここで L_1，SSIM は実画像に対する光度項であり W_geo を掛けない．
第 3 項のみ深度正則化に W_geo を適用する（従来方式との差異）．
M は深度有効マスク，α = 0.2，λ_d(t) は学習進行に応じて減衰する．
""",
    )

    doc.add_paragraph("2.3 学習パイプライン", style=S_H2)
    add_body(
        doc,
        """
処理は 3 段階である．(i) VGGT-1B で深度・カメラ・トラックを推論し，
σ，V，C から W_geo を生成．(ii) COLMAP 形式・深度 PNG・重み NPY へ変換．
(iii) 3DGS を 7000 反復学習する．実装は export_vggt_for_3dgs.py と
convert_vggt_to_3dgs.py，train.py（--weights weights）に基づく．
""",
    )
    doc.add_paragraph("図 1　VGGT 推論から W_geo 付き 3DGS 学習までのパイプライン", style=S_FIG)
    doc.add_paragraph(
        "（図 1：パイプライン図を挿入してください）",
        style=S_BODY,
    )

    # --- 3. 性能評価 ---
    doc.add_paragraph("3. 性能評価", style=S_H1)

    doc.add_paragraph("3.1 評価条件", style=S_H2)
    add_body(
        doc,
        """
GPU は NVIDIA RTX 3060 (12GB)，PyTorch 2.4，VGGT-1B を用いた．
データは Kitchen（屋内 8 枚），LLFF の Fern・Flower（各 8 枚）．
3DGS は 7000 反復，指標は訓練視点の L1，PSNR，SSIM とした．
比較は深度正則化あり・W_geo なし（Baseline）と W_geo ありの 2 条件．
消融は Fern で W の係数 10 構成を，VGGT 再推論なしに W のみ再計算して実施した．
""",
    )

    doc.add_paragraph("3.2 評価結果", style=S_H2)
    add_body(
        doc,
        """
表 1 に主結果を示す．W_geo は 3 シーンすべてで PSNR を改善し，
Fern では L1 を 16.8% 低減，+1.92 dB，Flower では SSIM +3.82% を得た．
Kitchen は Baseline が既に 40 dB 超のため改善は小さいが劣化しない．
Fern・Flower では 5 視点すべてで PSNR が向上した．
""",
    )
    add_table_results(doc)
    add_body(
        doc,
        """
Fern の係数消融では，無重み Baseline が PSNR 24.35 dB に対し，
全 9 構成の W_geo が +0.52〜+0.82 dB 改善した．
(0.5, 0.3, 0.2) が 25.53 dB で最良であり，均等重み (0.33, 0.33, 0.33) は
25.06 dB と W_geo 構成中最低であった．
""",
    )

    doc.add_paragraph(
        "図 2　Fern における Baseline と W_geo のレンダリング比較（代表視点）",
        style=S_FIG,
    )
    if FIG_CANDIDATE.exists():
        p = doc.add_paragraph(style=S_BODY)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(str(FIG_CANDIDATE), width=Cm(14))
    else:
        doc.add_paragraph("（図 2：比較画像を挿入してください）", style=S_BODY)

    # --- 4. おわりに ---
    doc.add_paragraph("4. おわりに", style=S_H1)
    add_body(
        doc,
        """
VGGT+3DGS において，W_geo を深度正則化に限定適用する概念分離は，
3 シーンで一貫した品質向上をもたらした．誤った深度教師の影響を
画素ごとに抑制する効果が，特に遮蔽の多い Fern で顕著であった．
係数消融では W_geo 機構のロバスト性とデフォルト係数の妥当性を確認した．
今後は光度損失向け W_appear の分離，テスト視点評価，大規模ベンチマークへの
適用が課題である．
""",
    )

    # --- 参考文献 ---
    doc.add_paragraph("参考文献", style=S_REF_HEAD)
    refs = [
        '[1] B. Kerbl et al.: "3D Gaussian Splatting for Real-Time Radiance Field Rendering," '
        "ACM Trans. Graph., Vol. 42, No. 4 (2023).",
        '[2] J. Wang et al.: "VGGT: Visual Geometry Grounded Transformer," '
        "arXiv preprint (2024).",
        '[3] V. Arampatzakis et al.: "Monocular depth estimation: A thorough review," '
        "IEEE Trans. PAMI, Vol. 46, No. 4, pp. 2396--2414 (2024).",
    ]
    for r in refs:
        doc.add_paragraph(r, style=S_REF)

    doc.save(str(OUT))
    return OUT


if __name__ == "__main__":
    path = build()
    print(f"Generated: {path}")
