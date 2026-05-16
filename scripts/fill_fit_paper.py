#!/usr/bin/env python3
"""Copy FITpaper.docx and fill with draft content (preserve styles)."""
from pathlib import Path
import shutil

from docx import Document

ROOT = Path(__file__).resolve().parents[1]
TEMPLATE = Path(r"d:\Downloads\FITpaper.docx")
OUT = ROOT / "FIT_paper_draft.docx"

# Ordered replacements: match template placeholder fragments (unique enough)
REPLACEMENTS = [
    ("FIT 論文テンプレート", "VGGT 誘導 3D Gaussian Splatting における\n幾何信頼性重み W_geo による深度正則化"),
    ("Format of FIT Paper", "W_geo-Weighted Depth Regularization for VGGT-Guided 3D Gaussian Splatting"),
    ("岡本 みかげ　　奥山 直美", "郭　某某"),
    ("Mikage Okamoto　 Naomi Okuyama", "Guo Dedi"),
]

# First long FIT0 block after figure -> 概要 body (match start of dummy text)
ABSTRACT_JA = (
    "Visual Geometry Grounded Transformer (VGGT) の幾何出力を3D Gaussian Splatting (3DGS) の学習に統合する際，"
    "従来手法は単一の重みマップ W を光度損失と深度正則化損失の双方に適用しており，"
    "「VGGT の低信頼領域」と「3DGS の再構築困難領域」という異なる概念を混同している．"
    "本稿では，幾何信頼性 W_geo を深度正則化損失にのみ適用する概念分離を提案する．"
    "W_geo は深度不確実性，可視性，幾何一致性の加重平均 (0.5, 0.3, 0.2) で構成する．"
    "3 シーン・各 8 視点・7000 反復の実験で，W_geo は全シーンで PSNR を改善し（最大 +1.92 dB），"
    "10 構成の消融でデフォルト係数の妥当性を確認した．"
)

DUMMY_PREFIX = "この文章はダミーです"


def replace_in_paragraph(paragraph, old: str, new: str) -> bool:
    if old not in paragraph.text:
        return False
    # Preserve first run style when possible
    if paragraph.runs:
        paragraph.runs[0].text = paragraph.text.replace(old, new)
        for run in paragraph.runs[1:]:
            run.text = ""
    else:
        paragraph.text = paragraph.text.replace(old, new)
    return True


def main():
    if not TEMPLATE.exists():
        raise SystemExit(f"Template not found: {TEMPLATE}")
    shutil.copy2(TEMPLATE, OUT)
    doc = Document(str(OUT))

    for old, new in REPLACEMENTS:
        for p in doc.paragraphs:
            replace_in_paragraph(p, old, new)

    filled_abstract = False
    for p in doc.paragraphs:
        if not filled_abstract and DUMMY_PREFIX in p.text and p.style and "FIT0" in (p.style.name or ""):
            if p.runs:
                p.runs[0].text = ABSTRACT_JA
                for run in p.runs[1:]:
                    run.text = ""
            else:
                p.text = ABSTRACT_JA
            filled_abstract = True
            break

    doc.save(str(OUT))
    print(f"Wrote: {OUT}")


if __name__ == "__main__":
    main()
